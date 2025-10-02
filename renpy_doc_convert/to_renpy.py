import logging
import re
from collections import namedtuple

from docx.text.run import Run
from docx.shared import RGBColor
from docx.document import Document

from renpy_doc_convert.consolidate import TextChunk, TextType, ConsolidateTextType
from typing import List, Dict, Optional, Tuple

INDENTATION_SPACES = 2
DEFAULT_FONT_SIZE = 11.0
DEFAULT_FONT_COLOR = "000000" # Hexadecimal Black

# name_found : bool
# text : string
CharNameReturn = namedtuple('CharNameReturn', ['name_found', 'text'])

class CharacterDefinition:
  def __init__(self, short_name: str, full_name: str, color: str):
    self.short_name = short_name
    self.full_name = full_name
    self.color = color

class ConvertToRenpy:

  def __init__(self, document: Document, chunks: List[TextChunk], output_file_path: str):
    self.chunks: List[TextChunk] = chunks
    self.output_file_path: str = output_file_path
    self.font_standards: FontStandards = FontStandards(document, chunks)
    self.renpy_styler = RenpyStyling(self.font_standards)
    self.character_definitions: Dict[str, CharacterDefinition] = {}
    self.use_character_definitions = False
    
    logging.debug("Finish with initializing ConvertToRenpy constructor")

  def get_label(self, output_file_path: str) -> str:
    path_list: List[str] = output_file_path.split('/')
    
    filename_with_extension: str = path_list[len(path_list) - 1]
    label = filename_with_extension.split(".")[0]

    logging.debug("Renpy label: {0}".format(label))

    return label

  def parse_character_definitions(self) -> Tuple[bool, int]:
    """
    Parse character definitions from the documents.
    Returns (found, end_index) where end_index is the chunk after Characters{}
    """
    if not self.chunks:
      return False, 0
    
    # Check if first chunk is CHARACTER_DEF type
    if self.chunks[0].text_type != TextType.CHARACTER_DEF:
      return False, 0
    
    # Parse the character definition chunk
    for paragraph in self.chunks[0].paragraphs:
      text = paragraph.text.strip()
      
      if text.startswith("Characters{") or text == "}":
        continue
      
      # Parse character line
      if "=" in text:
        parts = text.split("=", 1)
        short_name = parts[0].strip()
        full_name = parts[1].strip().rstrip(',')
        
        # Extract color - look for colored text in the ENTIRE paragraph
        # The color is usually applied to the character name part
        color = None
        for run in paragraph.runs:
          # Check if this run has the short name (like "E" or "F")
          if short_name in run.text:
            if run.font.color and run.font.color.rgb:
              color_hex = "#{0}".format(run.font.color.rgb)
              # Only use non-black colors
              if color_hex != "#000000":
                color = color_hex
                break
        
        # If no color found in short name, check the full name part
        if not color:
          for run in paragraph.runs:
            if full_name.replace(',', '') in run.text:
              if run.font.color and run.font.color.rgb:
                color_hex = "#{0}".format(run.font.color.rgb)
                if color_hex != "#000000":
                  color = color_hex
                  break
        
        # If still no color, get any non-black color from the line
        if not color:
          color = self.extract_color_from_paragraph(paragraph, short_name)
        
        self.character_definitions[short_name] = CharacterDefinition(
          short_name=short_name,
          full_name=full_name,
          color=color if color else "#FFFFFF"
        )
    
    if self.character_definitions:
      self.use_character_definitions = True
      logging.debug(f"Found {len(self.character_definitions)} character definitions")
      return True, 1
    
    return False, 0

  def extract_color_from_paragraph(self, paragraph, target_text: str) -> Optional[str]:
    """Extract color from a paragraph's runs - look for any colored text in the line"""
    for run in paragraph.runs:
      # Look for any colored text in this paragraph (the character name is usually colored)
      if run.font.color and run.font.color.rgb:
        color_hex = "#{0}".format(run.font.color.rgb)
        # Skip if it's just black text
        if color_hex != "#000000" and color_hex != "#FFFFFF":
          return color_hex
    return None

  def get_chunk_full_text(self, chunk: TextChunk) -> str:
    """Get the full text of a chunk"""
    text = ""
    for paragraph in chunk.paragraphs:
      text += paragraph.text
    return text.strip()

  def is_comment_line(self, chunk: TextChunk) -> Tuple[bool, str]:
    """Check if chunk is a comment and return (is_comment, comment_text)"""
    text = self.get_chunk_full_text(chunk).strip()
    
    # Check for # comment
    if text.startswith("#"):
      return True, text
    
    # Check for () comment
    if text.startswith("(") and text.endswith(")"):
      comment_text = text[1:-1].strip()
      return True, f"# {comment_text}"
    
    return False, ""

  def is_label_marker(self, chunk: TextChunk) -> Tuple[bool, str]:
    """Check if chunk is a label marker like '== label_name =='"""
    text = self.get_chunk_full_text(chunk).strip()
    
    match = re.match(r'^==\s*([\w_]+)\s*==$', text)
    if match:
      return True, match.group(1)
    
    return False, ""

  def is_menu_choice(self, chunk: TextChunk) -> Tuple[bool, str, str]:
    """Check if chunk is a menu choice and return (is_menu, choice_text, jump_label)"""
    text = self.get_chunk_full_text(chunk).strip()
    
    # Remove leading dash (both - and –)
    if text.startswith("-") or text.startswith("–"):
      text = text[1:].strip()
    elif text.startswith("    ") or text.startswith("\t"):
      text = text.lstrip()
      if text.startswith("-") or text.startswith("–"):
        text = text[1:].strip()
    else:
      return False, "", ""
    
    # Check for == label pattern
    if "==" in text:
      parts = text.split("==")
      choice_text = parts[0].strip()
      jump_label = parts[1].strip() if len(parts) > 1 else ""
      return True, choice_text, jump_label
    
    return True, text, ""

  def extract_character_styling(self, chunk: TextChunk) -> str:
      """Extract styling applied to the character name portion before the colon"""
      if chunk.text_type != TextType.DIALOGUE:
        return ""
      
      character_styles = []
      found_colon = False
      
      for paragraph in chunk.paragraphs:
        for run in paragraph.runs:
          if not found_colon:
            if ":" in run.text:
              # This run has the colon - extract the character name part
              char_part = run.text.split(":", maxsplit=1)[0]
              # Apply styling to just the character name
              styled_name = self.renpy_styler.apply_styling_to_text(char_part, run)
              character_styles.append(styled_name)
              found_colon = True
              break
            else:
              # This whole run is part of the character name
              styled_name = self.renpy_styler.apply_styling_to_text(run.text, run)
              character_styles.append(styled_name)
        
        if found_colon:
          break
      
      # If we collected multiple runs for the character name
      result = "".join(character_styles)
      
      # Return the styled character name, or empty if same as plain character name
      return result if result else "" 
  def output_renpy_text(self):
    logging.debug("Output renpy text to file")
    with open(self.output_file_path, "w", encoding="utf-8") as file:
      # Parse character definitions first
      has_char_defs, skip_until = self.parse_character_definitions()
      
      if has_char_defs:
        # Write character definitions at the very top
        for char_name, char_def in self.character_definitions.items():
          file.write(f'define {char_def.short_name} = Character("{char_def.full_name}", color="{char_def.color}")\n')
        file.write("\n")
      
      # Write label after definitions
      label = self.get_label(self.output_file_path)
      file.write("label {0}:\n\n".format(label))
      
      # Process chunks
      logging.debug("Processing {0} text chunk(s)".format(len(self.chunks)))
      
      start_idx = skip_until if has_char_defs else 0
      i = start_idx
      
      while i < len(self.chunks):
        chunk = self.chunks[i]
        
        # Check for label marker
        is_label, label_name = self.is_label_marker(chunk)
        if is_label:
          file.write(f"label {label_name}:\n")
          i += 1
          continue
        
        # Check for comment
        is_comment, comment_text = self.is_comment_line(chunk)
        if is_comment:
          file.write(f"{comment_text}\n")
          i += 1
          continue
        
        # Check if this is a dialogue line followed by menu choices
        if chunk.text_type == TextType.DIALOGUE and i + 1 < len(self.chunks):
          next_is_menu, _, _ = self.is_menu_choice(self.chunks[i + 1])
          
          if next_is_menu:
            # This is a menu prompt line
            text = self.handle_styling(chunk)
            text = self.handle_escape_characters(text)
            
            # Write the prompt
            if self.use_character_definitions and chunk.character:
              char_short = chunk.character
              if char_short in self.character_definitions:
                file.write(f'  {char_short} "{text}"\n')
              else:
                file.write(f'  "{chunk.character}" "{text}"\n')
            else:
              if chunk.character:
                file.write(f'  "{chunk.character}" "{text}"\n')
              else:
                file.write(f'  "{text}"\n')
            
            # Write menu
            file.write("  menu:\n")
            
            # Process menu choices
            i += 1
            while i < len(self.chunks):
              is_menu_item, choice_text, jump_label = self.is_menu_choice(self.chunks[i])
              if not is_menu_item:
                break
              
              file.write(f'    "{choice_text}":\n')
              if jump_label:
                file.write(f'      jump {jump_label}\n')
              i += 1
            
            file.write("\n")
            continue
        
        # Regular dialogue or narration
        text = self.handle_styling(chunk)
        text = self.handle_escape_characters(text)
        formatted_text = self.format_indentation(chunk, text)
        file.write(formatted_text)
        
        i += 1

  def handle_styling(self, chunk: TextChunk) -> str:
    text = ""

    for index, paragraph in enumerate(chunk.paragraphs):
      found_colon = False
      
      for run in paragraph.runs:
        run_text = run.text
        
        # For dialogue, skip everything until after the colon
        if chunk.text_type == TextType.DIALOGUE and not found_colon:
          if ":" in run_text:
            # This run contains the colon, take only the part after it
            run_text = run_text.split(":", maxsplit=1)[1].lstrip()
            found_colon = True
            
            # If nothing left after colon in this run, continue to next run
            if not run_text:
              continue
          else:
            # Still in character name part, skip this run
            continue

        # Process styling for the dialogue text
        original_text = run.text
        run.text = run_text
        appendtext = self.renpy_styler.process_run_for_styling(run)
        run.text = original_text
        text = text + appendtext

      if len(chunk.paragraphs) - 1 != index:
        text = text + "\n"
    
    return text

  def handle_escape_characters(self, text: str):
    """
    In renpy, there are special characters that need to be handled
    https://www.renpy.org/doc/html/text.html#escape-characters
    """
    if "\\" in text:
      text = text.replace("\\", "\\\\")

    if "\"" in text:
      text = text.replace('\"', '\\"')

    if "\'" in text:
      text = text.replace("\'", "\\'")

    if "%" in text:
      text = text.replace("%", "\\%")

    return text

  def format_indentation(self, chunk: TextChunk, text: str) -> str:
    if chunk.text_type == TextType.DIALOGUE:
      return self.format_dialogue(chunk, text)
    else:
      return self.format_non_dialogue(text)

  def format_dialogue(self, chunk: TextChunk, text: str) -> str:
    if self.use_character_definitions and chunk.character:
      # Use short character name if it exists in definitions
      char_short = chunk.character
      if char_short in self.character_definitions:
        character = char_short
        text_formatted = '"{0}"'.format(text)
        character_text = "{0} {1}".format(character, text_formatted)
      else:
        # Character not in definitions, extract any styling from character name
        styled_char = self.extract_character_styling(chunk)
        if styled_char and styled_char != chunk.character:
          # There's styling on the character name
          character = '"{0}"'.format(styled_char)
        else:
          character = '"{0}"'.format(chunk.character)
        text_formatted = '"{0}"'.format(text)
        character_text = "{0} {1}".format(character, text_formatted)
    else:
      # No character definitions, check for styling
      styled_char = self.extract_character_styling(chunk)
      if styled_char and styled_char != chunk.character:
        character = '"{0}"'.format(styled_char)
      else:
        character = '"{0}"'.format(chunk.character)
      text_formatted = '"{0}"'.format(text)
      character_text = "{0} {1}".format(character, text_formatted)

    # Add spaces to front of text
    indented_text = character_text.rjust(len(character_text) + INDENTATION_SPACES)

    # Add newline
    fulltext = "{0}\n".format(indented_text)

    return fulltext

  def format_non_dialogue(self, text: str):
    return "  \"" + text + "\"\n"

  def remove_character_name_in_text(self, text: str) -> CharNameReturn:
    if ":" in text:
      parts = text.split(":", maxsplit=1)
      remaining_text = parts[1].lstrip() if len(parts) > 1 else ""
      return CharNameReturn(name_found=True, text=remaining_text)
    else:
      return CharNameReturn(name_found=False, text=text)


class RenpyStyling:

  def __init__(self, font_stds):
    self.font_stds: FontStandards = font_stds

  def apply_styling_to_text(self, text: str, run: Run) -> str:
    """Apply styling to a specific text using a run's properties"""
    if run.bold:
      text = self.convert_bold(text)

    if run.italic:
      text = self.convert_italics(text)

    if run.underline:
      text = self.convert_underline(text)

    if run.font.size:
      text = self.convert_font_size(text, run.font.size.pt)

    if run.font.color and run.font.color.rgb:
      text = self.convert_font_color(text, run.font.color.rgb)

    if run.font.strike:
      text = self.convert_strike(text)
    
    return text

  def process_run_for_styling(self, run: Run) -> str:
    """Process styling for a run using its text"""
    return self.apply_styling_to_text(run.text, run)
  
  def convert_bold(self, text: str) -> str:
    return "{{b}}{0}{{/b}}".format(text)

  def convert_italics(self, text: str) -> str:
    return "{{i}}{0}{{/i}}".format(text)

  def convert_underline(self, text: str) -> str:
    return "{{u}}{0}{{/u}}".format(text)

  def convert_font_size(self, text: str, run_font_size: int) -> str:
    if self.font_stds.size == run_font_size:
      return text

    diff = run_font_size - self.font_stds.size
    diff = int(diff)

    return "{{size=+{0}}}{1}{{/size}}".format(diff, text)

  def convert_font_color(self, text: str, color: RGBColor):
    if self.font_stds.color == color:
      return text

    rgb_str: str = "#{0}".format(color)

    return "{{color={0}}}{1}{{/color}}".format(rgb_str, text)
    
  def convert_strike(self, text):
    return "{{s}}{0}{{/s}}".format(text)


class FontStandards:
  def __init__(self, document: Document, chunks: TextChunk):
    self.document = document
    self.chunks = chunks
    self.size = self.get_standard_font_size()
    self.color: RGBColor = self.get_standard_font_color()

    logging.debug("Font Size Standard: {0}".format(self.size))
    logging.debug("Font Color standard: {0}".format(self.color))
    logging.debug("Finish with initializing FontStandards constructor")

  def _get_size_first_line(self) -> int:
    if (len(self.chunks) and 
       len(self.chunks[0].paragraphs) and
       len(self.chunks[0].paragraphs[0].runs)):
      run = self.chunks[0].paragraphs[0].runs[0]
      
      if run.font.size:
        return run.font.size.pt

    return -1

  def _get_document_default(self) -> int:
    if (self.document != None and
       self.document.styles != None and
       self.document.styles.element != None):

      styles_elem = self.document.styles.element

      default_LXML = styles_elem.xpath('w:docDefaults/w:rPrDefault')
      if (len(default_LXML) != 0 and 
        len(default_LXML[0]) != 0):

        # https://github.com/python-openxml/python-docx/blob/master/docx/oxml/text/font.py#L52
        run = default_LXML[0][0] # Should be a type CT_RPr

        if run.sz_val:
          return run.sz_val.pt

    return -1

  def get_standard_font_size(self) -> int:
    """
    We have to get the standard font size for this document.

    1. Check if the first full line text has a font size. Use that as the priority.
    2. If the first fully line text has no font size, we check the document default
    3. If the document default has no font size, we use hard coded value
    """
    font_size = -1
    
    # Use font from first line
    if font_size == -1:
      font_size = self._get_size_first_line()

    # If font from first line cannot be found, use document default
    if font_size == -1:
      font_size = self._get_document_default()

    # If document default is not found, then use hard coded default font
    if font_size != -1:
      return font_size

    logging.info(
      "Found no font size in document." 
      "Using font size 11.0"
    )
    return DEFAULT_FONT_SIZE 

  def get_standard_font_color(self) -> RGBColor:
    """
    Get standard font color for this document.
    Assume standard font color is black
    """
    return RGBColor.from_string(DEFAULT_FONT_COLOR)