#doc-to-renpy/renpy_doc_convert/consolidate.py
from docx.document import Document
from docx.text.paragraph import Paragraph

from enum import Enum
from typing import List
import logging
import re

class TextType(Enum):
  DIALOGUE = 1
  NARRATION = 2
  SOUND = 3
  NONE = 4
  COMMENT = 5
  LABEL_MARKER = 6
  MENU_CHOICE = 7
  CHARACTER_DEF = 8

class ConsolidateTextType(Enum):
  """Additional text types for special Ren'Py constructs"""
  COMMENT = 5
  LABEL_MARKER = 6
  MENU_CHOICE = 7
  CHARACTER_DEF = 8

class TextChunk:

  def __init__(self):
    self.paragraphs: List[Paragraph] = []
    self.text_type: TextType = TextType.NONE
    self.character: str = ""

class Consolidate:
  
  def __init__(self, document: Document):
    self.document = document
    self.text_chunks: list[TextChunk] = []
    self.doc_paragraphs: list[Paragraph] = document.paragraphs

    logging.debug("Finish with Consolidate constructor")
  
  def consolidate_paragraphs(self):
    in_character_block = False
    character_block_chunk = None

    logging.debug("Processing {0} paragraphs".format(len(self.doc_paragraphs)))

    for paragraph in self.doc_paragraphs:
      text = paragraph.text.strip()
      
      # Skip empty lines
      if text == "":
        continue
      
      # Check for Characters{ block start
      if text.startswith("Characters{"):
        in_character_block = True
        character_block_chunk = TextChunk()
        character_block_chunk.paragraphs.append(paragraph)
        character_block_chunk.text_type = TextType.CHARACTER_DEF
        self.text_chunks.append(character_block_chunk)
        continue
      
      # Inside character definition block
      if in_character_block:
        character_block_chunk.paragraphs.append(paragraph)
        if "}" in text:
          in_character_block = False
          character_block_chunk = None
        continue
      
      # Check for comment lines (# or ())
      if self.is_comment_line(text):
        chunk = TextChunk()
        chunk.paragraphs.append(paragraph)
        chunk.text_type = TextType.COMMENT
        self.text_chunks.append(chunk)
        continue
      
      # Check for label markers (== label ==)
      if self.is_label_marker(text):
        chunk = TextChunk()
        chunk.paragraphs.append(paragraph)
        chunk.text_type = TextType.LABEL_MARKER
        self.text_chunks.append(chunk)
        continue
      
      # Check for menu choices (indented lines with - or –)
      if self.is_menu_choice(text):
        chunk = TextChunk()
        chunk.paragraphs.append(paragraph)
        chunk.text_type = TextType.MENU_CHOICE
        self.text_chunks.append(chunk)
        continue
      
      # Regular text processing - each paragraph is its own chunk
      chunk = TextChunk()
      chunk.paragraphs.append(paragraph)
      chunk.text_type = self.get_text_type(paragraph)
      chunk.character = self.get_character(paragraph, chunk.text_type)
      self.text_chunks.append(chunk)

  def is_comment_line(self, text: str) -> bool:
    """Check if line is a comment (starts with # or wrapped in ())"""
    text = text.strip()
    return text.startswith("#") or (text.startswith("(") and text.endswith(")"))
  
  def is_label_marker(self, text: str) -> bool:
    """Check if line is a label marker like '== label_name =='"""
    text = text.strip()
    # Remove spaces around == for matching
    return bool(re.match(r'^==\s*[\w_]+\s*==$', text))
  
  def is_menu_choice(self, text: str) -> bool:
    """Check if line is a menu choice (starts with dash or indent)"""
    # Check for both regular dash (-) and en-dash (–)
    return (text.startswith("-") or text.startswith("–") or 
            text.startswith("    -") or text.startswith("    –") or
            text.startswith("\t-") or text.startswith("\t–"))

  def get_character(self, paragraph: Paragraph, text_type: TextType) -> str:
    if text_type == TextType.DIALOGUE:
      text = paragraph.text.strip()
      if ":" in text:
        return text.split(":", maxsplit=1)[0].strip()
    return ""
      
  def get_text_type(self, paragraph: Paragraph) -> TextType:
    text = paragraph.text.strip()

    if self.is_dialogue(text):
      return TextType.DIALOGUE
    elif self.is_sound(text):
      return TextType.SOUND
    elif self.is_narration(text):
      return TextType.NARRATION
    else:
      return TextType.NONE

  def is_dialogue(self, text: str) -> bool:
    """Check if text contains dialogue (has " : " but not at start for labels)"""
    if ":" not in text:
      return False
    
    # Make sure it's not a label marker or comment
    stripped = text.strip()
    if stripped.startswith("==") or stripped.startswith("#") or stripped.startswith("("):
      return False
    
    return True
    
  def is_sound(self, text: str) -> bool:
    return "*" in text

  def is_narration(self, text: str) -> bool:
    return ":" not in text